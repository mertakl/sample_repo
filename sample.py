from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def create_valid_eureka_docx(filename="test_document.docx", language="fr"):
    """
    Creates a DOCX document that will pass the update_titles_and_depths_eureka_nota function.
    
    Args:
        filename (str): Name of the output DOCX file
        language (str): Language for table of contents ("fr" or "nl")
    """
    # Create a new document
    doc = Document()
    
    # Define table of contents text based on language
    toc_text = {
        "fr": "table des matières",
        "nl": "inhoudsopgave"
    }
    
    # Add main title (will be detected as depth=0 via Subtitle style)
    main_title = doc.add_paragraph("Guide Principal de Documentation")
    main_title.style = "Subtitle"  # This creates depth=0 in the XML
    
    # Add main subtitle (also depth=0)
    subtitle = doc.add_paragraph("Manuel d'utilisation complet")
    subtitle.style = "Subtitle"
    
    # Add table of contents paragraph (required to pass validation)
    toc_paragraph = doc.add_paragraph(toc_text.get(language, "table des matières"))
    toc_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # Add some regular text
    doc.add_paragraph("Ce document présente les informations essentielles pour l'utilisation du système.")
    
    # Add Heading 1 (will be detected as depth=0 in XML due to Heading1 style)
    heading1 = doc.add_paragraph("Introduction Générale")
    heading1.style = "Heading 1"
    
    # Add some content under heading 1
    doc.add_paragraph("Cette section présente une vue d'ensemble du système et de ses fonctionnalités principales.")
    
    # Add Heading 2 (will be detected as depth=1)
    heading2 = doc.add_paragraph("Objectifs du Système")
    heading2.style = "Heading 2"
    
    # Add content
    doc.add_paragraph("Le système vise à améliorer l'efficacité des processus documentaires.")
    
    # Add another Heading 2
    heading2_bis = doc.add_paragraph("Architecture Technique")
    heading2_bis.style = "Heading 2"
    
    # Add content
    doc.add_paragraph("L'architecture repose sur des composants modulaires et scalables.")
    
    # Add Heading 3 (will be detected as depth=2)
    heading3 = doc.add_paragraph("Composants Principaux")
    heading3.style = "Heading 3"
    
    # Add content
    doc.add_paragraph("Les composants incluent le parser, le validateur et l'interface utilisateur.")
    
    # Add another Heading 1 for variety
    heading1_bis = doc.add_paragraph("Configuration et Installation")
    heading1_bis.style = "Heading 1"
    
    # Add content
    doc.add_paragraph("Cette section détaille les étapes d'installation et de configuration.")
    
    # Add Heading 2 under the second main section
    heading2_config = doc.add_paragraph("Prérequis Système")
    heading2_config.style = "Heading 2"
    
    # Add content
    doc.add_paragraph("Le système nécessite Python 3.8+ et les dépendances listées ci-après.")
    
    # Add a table with caption (to test TableBlock handling)
    table = doc.add_table(rows=3, cols=2)
    table.style = 'Table Grid'
    
    # Add table headers
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Composant'
    hdr_cells[1].text = 'Version'
    
    # Add table data
    row_cells = table.rows[1].cells
    row_cells[0].text = 'Python'
    row_cells[1].text = '3.8+'
    
    row_cells = table.rows[2].cells
    row_cells[0].text = 'docx'
    row_cells[1].text = '0.8.11+'
    
    # Add a paragraph after table to serve as caption
    caption = doc.add_paragraph("Tableau 1: Composants et versions requises")
    caption.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # Add final content section
    heading2_final = doc.add_paragraph("Utilisation Avancée")
    heading2_final.style = "Heading 2"
    
    doc.add_paragraph("Pour une utilisation avancée, consultez la documentation technique complète.")
    
    # Save the document
    doc.save(filename)
    print(f"Document DOCX créé avec succès: {filename}")
    print("Ce document contient:")
    print("- Un titre principal et sous-titre (depth=0)")
    print("- Une table des matières")
    print("- Des titres hiérarchiques (Heading 1, 2, 3)")
    print("- Du contenu textuel")
    print("- Un tableau avec légende")
    
    return filename

def create_dutch_version():
    """Create a Dutch version of the document."""
    doc = Document()
    
    # Add main title and subtitle
    main_title = doc.add_paragraph("Hoofdgids voor Documentatie")
    main_title.style = "Subtitle"
    
    subtitle = doc.add_paragraph("Volledige gebruikershandleiding")
    subtitle.style = "Subtitle"
    
    # Add table of contents
    toc_paragraph = doc.add_paragraph("inhoudsopgave")
    toc_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # Add content structure
    doc.add_paragraph("Dit document bevat essentiële informatie voor het gebruik van het systeem.")
    
    heading1 = doc.add_paragraph("Algemene Introductie")
    heading1.style = "Heading 1"
    
    doc.add_paragraph("Deze sectie biedt een overzicht van het systeem en zijn hoofdfunctionaliteiten.")
    
    heading2 = doc.add_paragraph("Systeemdoelstellingen")
    heading2.style = "Heading 2"
    
    doc.add_paragraph("Het systeem is bedoeld om de efficiëntie van documentprocessen te verbeteren.")
    
    heading3 = doc.add_paragraph("Hoofdcomponenten")
    heading3.style = "Heading 3"
    
    doc.add_paragraph("De componenten omvatten de parser, validator en gebruikersinterface.")
    
    filename = "test_document_nl.docx"
    doc.save(filename)
    print(f"Nederlands document gemaakt: {filename}")
    return filename

# Example usage
if __name__ == "__main__":
    # Create French version
    french_doc = create_valid_eureka_docx("test_document_fr.docx", "fr")
    
    # Create Dutch version  
    dutch_doc = create_dutch_version()
    
    print("\nBoth documents created successfully!")
    print("These documents should pass the update_titles_and_depths_eureka_nota function because they contain:")
    print("1. Proper heading hierarchy (Heading 1, 2, 3 styles)")
    print("2. Main titles with Subtitle style (depth=0)")
    print("3. Table of contents text in the specified language")
    print("4. Mixed content types (text, tables)")
    print("5. Proper XML structure for depth detection")
